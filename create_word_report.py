#!/usr/bin/env python3
"""
Script to convert Markdown report to Word document
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def add_heading(doc, text, level=1):
    """Add a heading with proper formatting"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with optional formatting"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p

def add_code_block(doc, code):
    """Add a code block with monospace font"""
    p = doc.add_paragraph(code)
    p.style = 'Intense Quote'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    return p

def process_markdown_to_docx(md_file, docx_file):
    """Convert Markdown file to Word document"""
    doc = Document()
    
    # Set up document properties
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue
        
        # Handle horizontal rules
        if line.strip() == '---':
            doc.add_page_break()
            i += 1
            continue
        
        # Handle headings
        if line.startswith('#'):
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                # Remove anchor links
                text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
                add_heading(doc, text, level=min(level, 3))
            i += 1
            continue
        
        # Handle bold text
        if '**' in line:
            line = re.sub(r'\*\*([^\*]+)\*\*', r'\1', line)
            if line.strip():
                p = doc.add_paragraph()
                # Split and format bold parts
                parts = re.split(r'(\*\*[^\*]+\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.strip():
                        p.add_run(part)
            i += 1
            continue
        
        # Handle list items
        if line.strip().startswith('-') or line.strip().startswith('*'):
            item_text = re.sub(r'^[\-\*]\s+', '', line.strip())
            item_text = re.sub(r'\*\*([^\*]+)\*\*', r'\1', item_text)
            item_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', item_text)
            p = doc.add_paragraph(item_text, style='List Bullet')
            i += 1
            continue
        
        # Handle numbered lists
        if re.match(r'^\d+\.', line.strip()):
            item_text = re.sub(r'^\d+\.\s+', '', line.strip())
            item_text = re.sub(r'\*\*([^\*]+)\*\*', r'\1', item_text)
            item_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', item_text)
            p = doc.add_paragraph(item_text, style='List Number')
            i += 1
            continue
        
        # Handle tables (basic support)
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            if table_lines:
                # Parse table
                rows = []
                for tline in table_lines:
                    if '---' in tline:  # Skip separator line
                        continue
                    cells = [cell.strip() for cell in tline.split('|')[1:-1]]
                    if cells:
                        rows.append(cells)
                
                if rows:
                    # Create table
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            table.rows[row_idx].cells[col_idx].text = cell_data
                    
                    doc.add_paragraph()  # Add space after table
            continue
        
        # Handle regular paragraphs
        if line.strip():
            # Remove markdown formatting
            clean_line = re.sub(r'\*\*([^\*]+)\*\*', r'\1', line)
            clean_line = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', clean_line)
            clean_line = re.sub(r'`([^`]+)`', r'\1', clean_line)
            
            if clean_line.strip():
                doc.add_paragraph(clean_line.strip())
        else:
            # Empty line - could indicate paragraph break
            if i > 0 and i < len(lines) - 1:
                # Only add space if between content
                pass
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✅ Word document created successfully: {docx_file}")
    print(f"📄 Document contains {len(doc.paragraphs)} paragraphs")
    print(f"📊 Document contains {len(doc.tables)} tables")

if __name__ == "__main__":
    md_file = "/Users/paarthsanyal/Desktop/sentiment/Sentiment_Analysis_Project_Report.md"
    docx_file = "/Users/paarthsanyal/Desktop/sentiment/Sentiment_Analysis_Project_Report.docx"
    
    process_markdown_to_docx(md_file, docx_file)
